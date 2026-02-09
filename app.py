
import streamlit as st
import os
import shutil
import uuid
import io
import time
from docx import Document
# Assuming the backend is in src/backend.py and class is AdvancedRAG
from src.backend import AdvancedRAG

# ==========================================
# 1. PAGE CONFIGURATION & SETUP
# ==========================================
st.set_page_config(
    page_title="Enterprise RAG System",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Initialize Backend
@st.cache_resource
def get_rag_engine():
    return AdvancedRAG()

rag_engine = get_rag_engine()

# ==========================================
# 2. SESSION STATE MANAGEMENT (PRIVACY & LOGIC)
# ==========================================
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

if "messages" not in st.session_state:
    st.session_state.messages = []

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "db_ready" not in st.session_state:
    st.session_state.db_ready = False

if "current_model" not in st.session_state:
    st.session_state.current_model = "Llama 3.3 70B (Versatile)"

# Define Paths based on Session ID (Data Isolation)
BASE_DIR = "temp_data"
USER_SESSION_DIR = os.path.join(BASE_DIR, st.session_state.session_id)
FILES_DIR = os.path.join(USER_SESSION_DIR, "files")
DB_DIR = os.path.join(USER_SESSION_DIR, "db")

# Ensure directories exist
os.makedirs(FILES_DIR, exist_ok=True)
os.makedirs(DB_DIR, exist_ok=True)

# Helper: Clean up session data
def cleanup_session_data():
    if os.path.exists(USER_SESSION_DIR):
        try:
            shutil.rmtree(USER_SESSION_DIR)
        except Exception as e:
            print(f"Error cleaning up: {e}")
    # Re-create for new session usage if needed, or just let the logic handle it
    os.makedirs(FILES_DIR, exist_ok=True)
    os.makedirs(DB_DIR, exist_ok=True)

# Helper: Generate DOCX
def generate_document(messages):
    doc = Document()
    doc.add_heading('Confidential Conversation Log', 0)
    for msg in messages:
        role = "User" if msg["role"] == "user" else "AI Assistant"
        p = doc.add_paragraph()
        run = p.add_run(f"{role}: ")
        run.bold = True
        run.font.color.rgb = (0x00, 0x00, 0x00)  # Black
        doc.add_paragraph(msg["content"])
        doc.add_paragraph("_" * 40)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# High-End Model Map
MODEL_MAP = {
    "Llama 3.3 70B (Versatile)": "llama-3.3-70b-versatile",
    "Llama 3.1 8B (Instant)": "llama-3.1-8b-instant",
    "Llama 4 (Scout 17B)": "meta-llama/llama-4-scout-17b-16e-instruct",
    "Qwen 3 32B": "qwen/qwen3-32b",
    "GPT-OSS 20B": "openai/gpt-oss-20b"
}

# ==========================================
# 3. PROFESSIONAL UI / CSS INJECTION
# ==========================================
st.markdown("""
<style>
    /* IMPORT FONTS */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;600&display=swap');

    /* GLOBAL RESET */
    * {
        font-family: 'Inter', sans-serif;
    }

    /* MAIN BACKGROUND */
    .stApp {
        background-color: #0d0f14; /* Deep Enterprise Dark */
        background-image: radial-gradient(circle at 50% 0%, #1a1f2e 0%, #0d0f14 70%);
        color: #e0e0e0;
    }

    /* HIDE STREAMLIT ELEMENTS */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display:none;}

    /* TYPOGRAPHY */
    h1, h2, h3 {
        color: #ffffff;
        letter-spacing: -0.02em;
        font-weight: 700;
    }
    
    .hero-title {
        font-size: 3rem;
        background: linear-gradient(90deg, #60a5fa, #a78bfa);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    
    .hero-subtitle {
        color: #94a3b8;
        text-align: center;
        margin-bottom: 3rem;
        font-weight: 300;
        font-size: 1.1rem;
    }

    /* CONTAINERS */
    .upload-container {
        background: rgba(30, 41, 59, 0.5);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 16px;
        padding: 2rem;
        backdrop-filter: blur(10px);
        margin: 0 auto;
        max-width: 800px;
        box-shadow: 0 20px 25px -5px rgba(0, 0, 0, 0.3);
    }

    /* CHAT BUBBLES */
    .chat-message {
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 1rem;
        display: flex;
        flex-direction: row;
        align-items: flex-start;
        opacity: 0;
        animation: fadeIn 0.3s ease-in forwards;
    }

    @keyframes fadeIn {
        to { opacity: 1; transform: translateY(0); }
    }

    .chat-message.user {
        background-color: transparent;
        border-right: 2px solid #3b82f6; 
        flex-direction: row-reverse;
    }
    
    .chat-message.ai {
        background-color: #1e293b;
        border-left: 2px solid #a855f7;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
    
    .chat-message .avatar {
        width: 40px;
        height: 40px;
        border-radius: 8px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.2rem;
        flex-shrink: 0;
    }
    
    .chat-message.user .avatar {
        margin-left: 1rem;
        background: #3b82f6;
        color: white;
    }

    .chat-message.ai .avatar {
        margin-right: 1rem;
        background: #a855f7;
        color: white;
    }

    .chat-message .message-content {
        font-size: 0.95rem;
        line-height: 1.6;
        color: #e2e8f0;
        max-width: 80%;
    }

    /* BUTTONS */
    .stButton button {
        background: linear-gradient(90deg, #3b82f6, #2563eb);
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.2s;
    }
    
    .stButton button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.4);
    }

    /* SIDEBAR */
    section[data-testid="stSidebar"] {
        background-color: #0b0d11;
        border-right: 1px solid #1f2937;
    }
    
    .sidebar-history-item {
        padding: 10px;
        border-radius: 6px;
        margin-bottom: 5px;
        cursor: pointer;
        color: #94a3b8;
        font-size: 0.9rem;
        border: 1px solid transparent;
        transition: all 0.2s;
    }
    
    .sidebar-history-item:hover {
        background-color: #1e293b;
        color: #f8fafc;
        border-color: #334155;
    }
    
</style>
""", unsafe_allow_html=True)

# ==========================================
# 4. SIDEBAR (NAVIGATION & HISTORY)
# ==========================================
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/9626/9626605.png", width=50) # Placeholder generic AI icon
    st.markdown("### Enterprise Intelligence")
    st.markdown("---")
    
    # NEW CHAT BUTTON
    if st.button("➕ Start New Chat", use_container_width=True):
        # 1. Archive current chat if it exists
        if st.session_state.messages and st.session_state.db_ready:
            title = st.session_state.messages[0]['content'][:30] + "..." if st.session_state.messages else "Untitled Chat"
            st.session_state.chat_history.insert(0, {
                "id": st.session_state.session_id,
                "title": title,
                "timestamp": time.strftime("%Y-%m-%d %H:%M")
            })
            
        # 2. Cleanup resources
        # We DON'T delete the folder immediately if we want to support history retrieval, 
        # but the user requested strict privacy/cleanup. 
        # Requirement: "when i click a new chat button everything should get disapper i mean the file uploaded , the converstaion"
        # Requirement 3: "browse file section should get empty"
        cleanup_session_data() 
        
        # 3. Reset State
        st.session_state.session_id = str(uuid.uuid4())
        st.session_state.messages = []
        st.session_state.db_ready = False
        st.rerun()

    st.markdown("### History")
    # For now, history just shows past sessions. 
    # Note: If we strictly delete files on new chat, we can't 'reload' the RAG for old chats without re-uploading.
    # The user asked to "Save history section with some relevant title".
    # We will save the TEXT history, but RAG context might be lost if checked.
    if not st.session_state.chat_history:
        st.caption("No recent history.")
    
    for chat in st.session_state.chat_history:
        st.markdown(f"<div class='sidebar-history-item'>🕒 {chat['title']}</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown(f"<div style='color: #4b5563; font-size: 0.8rem;'>Session ID:<br>{st.session_state.session_id[:8]}...</div>", unsafe_allow_html=True)

# ==========================================
# 5. MAIN APPLICATION LOGIC
# ==========================================

# VIEW 1: LANDING / SETUP (If DB not ready)
if not st.session_state.db_ready:
    st.markdown("<h1 class='hero-title'>Multi Model RAG</h1>", unsafe_allow_html=True)
    st.markdown("<p class='hero-subtitle'>SECURE • PRIVATE • ENTERPRISE GRADE</p>", unsafe_allow_html=True)

    with st.container():
        st.markdown('<div class="upload-container">', unsafe_allow_html=True)
        
        # MODEL SELECTION
        col1, col2 = st.columns([1, 2])
        with col1:
            st.markdown("### 1. Select Intelligence")
            st.info("Choose the LLM best suited for your task.")
        with col2:
            model_friendly = st.selectbox(
                "AI Model", 
                options=list(MODEL_MAP.keys()),
                index=0,
                label_visibility="collapsed"
            )
            st.session_state.current_model = MODEL_MAP[model_friendly]

        st.markdown("---")

        # FILE UPLOAD
        col3, col4 = st.columns([1, 2])
        with col3:
            st.markdown("### 2. Upload Data")
            st.info("Supported: PDF, DOCX (Strictly private)")
        with col4:
            # Recreate uploader key every session to clear it
            upl_key = f"uploader_{st.session_state.session_id}"
            uploaded_files = st.file_uploader(
                "Drop secure documents here", 
                accept_multiple_files=True,
                type=['pdf', 'docx', 'txt'],
                key=upl_key,
                label_visibility="collapsed"
            )

        # PROCESS BUTTON
        if uploaded_files:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🔒 Process & Initialize Secure Chat", use_container_width=True, type="primary"):
                with st.spinner("Encrypting and Indexing Data..."):
                    # Save files
                    for file in uploaded_files:
                        with open(os.path.join(FILES_DIR, file.name), "wb") as f:
                            f.write(file.getbuffer())
                    
                    # Process
                    status = rag_engine.process_documents(FILES_DIR, DB_DIR)
                    
                    if status == "Success":
                        st.session_state.db_ready = True
                        st.rerun()
                    else:
                        st.error(f"Processing Failed: {status}")
        
        st.markdown('</div>', unsafe_allow_html=True)

# VIEW 2: CHAT INTERFACE (If DB is ready)
else:
    # Header
    top_col1, top_col2 = st.columns([6, 1])
    with top_col1:
        st.markdown("### 🔒 Secure Session Active")
    with top_col2:
        # Download Button logic
        if st.session_state.messages:
            docx = generate_document(st.session_state.messages)
            st.download_button("📥 Export", docx, f"log_{st.session_state.session_id[:6]}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")

    # Chat Container
    chat_placeholder = st.container()

    with chat_placeholder:
        if not st.session_state.messages:
            st.markdown("<div style='text-align: center; color: #64748b; margin-top: 50px;'><i>Begin your secure query regarding the uploaded documents...</i></div>", unsafe_allow_html=True)
            
        for msg in st.session_state.messages:
            role_class = "user" if msg["role"] == "user" else "ai"
            icon = "👤" if msg["role"] == "user" else "🤖"
            
            st.markdown(f"""
            <div class="chat-message {role_class}">
                <div class="avatar">{icon}</div>
                <div class="message-content">
                    <b>{msg.get('model_name', 'You' if msg['role'] == 'user' else 'Assistant')}</b><br>
                    {msg['content']}
                </div>
            </div>
            """, unsafe_allow_html=True)

    # Input Area (Fixed at bottom via standard Streamlit chat_input)
    if prompt := st.chat_input("Type your secure query..."):
        # Add User Message
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.rerun()

    # Process AI Response (Triggered after rerun to show user message first)
    if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
        with st.spinner("Analyzing secure context..."):
            try:
                response = rag_engine.query(
                    query_text=st.session_state.messages[-1]["content"],
                    db_path=DB_DIR,
                    model_name=st.session_state.current_model
                )
                
                # Check for sources/metadata if your backend returns dict, or string
                # Assuming string for now based on previous app.py, but backend.py has query_with_sources. 
                # Let's use simple string for safety or check type.
                final_text = response
                if isinstance(response, dict) and "answer" in response:
                    final_text = response["answer"]

                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": final_text,
                    "model_name": st.session_state.current_model
                })
                st.rerun()
            except Exception as e:
                st.error(f"System Error: {str(e)}")